"""
Enhanced Document Processor combining original ModernDocumentProcessor with new intelligent chunking
"""

import hashlib
import logging
import asyncio
from pathlib import Path
from typing import Dict, List, Any, Optional
import json

# Original imports
from docling.document_converter import DocumentConverter
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.pipeline.standard_pdf_pipeline import StandardPdfPipeline
import pypdf
import docx
from unstructured.partition.auto import partition
import nltk
from sentence_transformers import SentenceTransformer

# New imports
import openai
from openai import OpenAI

from apps.api.core.config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Original document processor for backward compatibility"""
    
    def __init__(self):
        self.client = OpenAI(api_key=settings.openai_api_key)
        self.max_chunk_size = 1000
        self.overlap_size = 200
        
    async def process_document(self, file_path: str) -> List[Dict[str, Any]]:
        """Original process_document method - preserved exactly"""
        try:
            file_path = Path(file_path)
            
            # Extract text based on file type
            if file_path.suffix.lower() == '.pdf':
                text_content = await self._extract_pdf_text(file_path)
            elif file_path.suffix.lower() == '.docx':
                text_content = await self._extract_docx_text(file_path)
            elif file_path.suffix.lower() == '.txt':
                text_content = await self._extract_txt_text(file_path)
            else:
                # Use unstructured for other formats
                text_content = await self._extract_with_unstructured(file_path)
            
            # Intelligent chunking
            chunks = await self._intelligent_chunking(text_content)
            
            # Extract multi-level context
            processed_chunks = []
            for i, chunk in enumerate(chunks):
                chunk_data = {
                    "chunk_id": f"chunk_{i}",
                    "content": chunk,
                    "position": i,
                    "metadata": await self._extract_metadata(chunk, i, len(chunks)),
                    "local_context": await self._extract_local_context(chunk, chunks, i),
                    "document_context": await self._extract_document_context(chunk, text_content),
                    "content_hash": hashlib.md5(chunk.encode()).hexdigest()
                }
                processed_chunks.append(chunk_data)
            
            logger.info(f"✅ Processed document with {len(processed_chunks)} chunks")
            return processed_chunks
            
        except Exception as e:
            logger.error(f"❌ Document processing failed: {e}")
            raise
    
    # All original methods preserved exactly...
    async def _extract_pdf_text(self, file_path: Path) -> str:
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
            return text.strip()
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    async def _extract_txt_text(self, file_path: Path) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise
    
    async def _extract_with_unstructured(self, file_path: Path) -> str:
        try:
            elements = partition(filename=str(file_path))
            return "\n\n".join([str(element) for element in elements])
        except Exception as e:
            logger.error(f"Unstructured extraction failed: {e}")
            raise
    
    async def _intelligent_chunking(self, text: str) -> List[str]:
        try:
            paragraphs = text.split('\n\n')
            chunks = []
            current_chunk = ""
            
            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                
                if len(current_chunk) + len(paragraph) > self.max_chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        current_chunk = paragraph
                    else:
                        sentences = paragraph.split('. ')
                        for sentence in sentences:
                            if len(current_chunk) + len(sentence) > self.max_chunk_size:
                                if current_chunk:
                                    chunks.append(current_chunk.strip())
                                current_chunk = sentence
                            else:
                                current_chunk += sentence + ". "
                else:
                    current_chunk += "\n\n" + paragraph
            
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            return chunks
            
        except Exception as e:
            logger.error(f"Chunking failed: {e}")
            return [text]
    
    async def _extract_metadata(self, chunk: str, position: int, total_chunks: int) -> Dict[str, Any]:
        return {
            "position": position,
            "total_chunks": total_chunks,
            "word_count": len(chunk.split()),
            "char_count": len(chunk),
            "position_ratio": position / total_chunks if total_chunks > 0 else 0,
            "is_beginning": position < 3,
            "is_middle": 3 <= position < total_chunks - 3,
            "is_ending": position >= total_chunks - 3,
        }
    
    async def _extract_local_context(self, chunk: str, all_chunks: List[str], position: int) -> Dict[str, str]:
        context = {}
        
        if position > 0:
            context["previous"] = all_chunks[position - 1][-200:]
        
        if position < len(all_chunks) - 1:
            context["next"] = all_chunks[position + 1][:200]
        
        return context
    
    async def _extract_document_context(self, chunk: str, full_text: str) -> Dict[str, Any]:
        try:
            prompt = f"""
            Analyze this document chunk and provide context:
            
            Chunk: "{chunk[:500]}..."
            
            Document length: {len(full_text)} characters
            
            Please provide:
            1. Main topic/theme of this chunk
            2. Document type (report, manual, article, etc.)
            3. Relevant keywords
            4. Importance level (1-5)
            
            Respond in JSON format.
            """
            
            response = await asyncio.to_thread(
                self.client.chat.completions.create,
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=200,
                temperature=0.1
            )
            
            result = response.choices[0].message.content
            
            try:
                return json.loads(result)
            except:
                return {"error": "Failed to parse AI response", "raw_response": result}
                
        except Exception as e:
            logger.error(f"Document context extraction failed: {e}")
            return {"error": str(e)}


class ModernDocumentProcessor:
    """Enhanced document processor preserving all original functionality"""
    
    def __init__(self):
        # Original properties
        self.chunk_size = getattr(settings, 'chunk_size', 1000)
        self.chunk_overlap = getattr(settings, 'chunk_overlap', 200)
        
        # Initialize Docling for advanced PDF processing
        self.pdf_pipeline_options = PdfPipelineOptions()
        self.pdf_pipeline_options.do_ocr = True
        self.pdf_pipeline_options.do_table_structure = True
        
        # Initialize sentence splitter
        try:
            nltk.download('punkt', quiet=True)
        except:
            pass
    
    async def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Enhanced process document with original functionality preserved"""
        logger.info(f"Processing document: {file_path} (type: {file_type})")
        
        # Calculate file hash for deduplication (original functionality)
        file_hash = self._calculate_file_hash(file_path)
        
        # Process based on file type (enhanced)
        if file_type.startswith("application/pdf"):
            result = await self._process_pdf_advanced(file_path)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            result = await self._process_docx_advanced(file_path)
        else:
            result = await self._process_with_unstructured(file_path)
        
        # Add file hash (original functionality)
        result["file_hash"] = file_hash
        
        # Create semantic chunks (enhanced functionality)
        result["chunks"] = await self._create_semantic_chunks(result)
        
        return result
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate file hash for deduplication - original functionality"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    async def _process_pdf_advanced(self, file_path: str) -> Dict[str, Any]:
        """Enhanced PDF processing using Docling"""
        try:
            converter = DocumentConverter()
            result = converter.convert(file_path)
            
            # Extract structured content
            full_text = result.document.export_to_markdown()
            
            # Extract metadata
            metadata = {
                "page_count": len(result.document.pages) if hasattr(result.document, 'pages') else 1,
                "title": getattr(result.document, 'title', Path(file_path).stem),
                "processing_method": "docling_advanced"
            }
            
            return {
                "content": full_text,
                "metadata": metadata,
                "structure": self._extract_document_structure(result.document)
            }
            
        except Exception as e:
            logger.warning(f"Advanced PDF processing failed, falling back to basic: {e}")
            return await self._process_pdf_basic(file_path)
    
    async def _process_pdf_basic(self, file_path: str) -> Dict[str, Any]:
        """Fallback basic PDF processing"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
            
            return {
                "content": text.strip(),
                "metadata": {
                    "processing_method": "pypdf_basic",
                    "title": Path(file_path).stem
                },
                "structure": {"sections": []}
            }
        except Exception as e:
            logger.error(f"Basic PDF processing failed: {e}")
            raise
    
    async def _process_docx_advanced(self, file_path: str) -> Dict[str, Any]:
        """Enhanced DOCX processing"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            sections = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
                    
                    # Detect headings
                    if para.style.name.startswith('Heading'):
                        sections.append({
                            "title": para.text,
                            "level": int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                        })
            
            full_text = "\n\n".join(content_parts)
            
            return {
                "content": full_text,
                "metadata": {
                    "paragraph_count": len(content_parts),
                    "section_count": len(sections),
                    "processing_method": "docx_advanced",
                    "title": Path(file_path).stem
                },
                "structure": {"sections": sections}
            }
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise
    
    async def _process_with_unstructured(self, file_path: str) -> Dict[str, Any]:
        """Process with unstructured library"""
        try:
            elements = partition(filename=str(file_path))
            content_parts = []
            sections = []
            
            for element in elements:
                content_parts.append(str(element))
                
                # Extract sections from titles
                if hasattr(element, 'category') and element.category == 'Title':
                    sections.append({
                        "title": str(element),
                        "level": 1
                    })
            
            full_text = "\n\n".join(content_parts)
            
            return {
                "content": full_text,
                "metadata": {
                    "element_count": len(elements),
                    "processing_method": "unstructured",
                    "title": Path(file_path).stem
                },
                "structure": {"sections": sections}
            }
            
        except Exception as e:
            logger.error(f"Unstructured processing failed: {e}")
            raise
    
    def _extract_document_structure(self, document) -> Dict[str, Any]:
        """Extract document structure for enhanced processing"""
        structure = {"sections": []}
        
        try:
            # This would be implemented based on the specific document structure
            # For now, return basic structure
            structure["sections"] = []
        except Exception as e:
            logger.warning(f"Structure extraction failed: {e}")
        
        return structure
    
    async def _create_semantic_chunks(self, processing_result: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Create semantic chunks with enhanced metadata"""
        content = processing_result.get("content", "")
        metadata = processing_result.get("metadata", {})
        structure = processing_result.get("structure", {})
        
        # Use intelligent chunking similar to original but enhanced
        chunks = []
        
        # Split by paragraphs first
        paragraphs = content.split('\n\n')
        current_chunk = ""
        chunk_index = 0
        
        for para_idx, paragraph in enumerate(paragraphs):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # Check if adding this paragraph exceeds chunk size
            if len(current_chunk) + len(paragraph) > self.chunk_size:
                if current_chunk:
                    # Create chunk
                    chunk_data = {
                        "chunk_index": chunk_index,
                        "content": current_chunk.strip(),
                        "chunk_content": current_chunk.strip(),  # For compatibility
                        "metadata": {
                            "position": chunk_index,
                            "word_count": len(current_chunk.split()),
                            "char_count": len(current_chunk),
                            "section_path": self._get_section_path(chunk_index, structure),
                            "token_count": len(current_chunk.split())  # Approximate
                        },
                        "contexts": {
                            "local": self._extract_local_context_enhanced(current_chunk, paragraphs, para_idx),
                            "document": metadata,
                            "global": {"document_title": metadata.get("title", "")},
                            "semantic": await self._extract_semantic_context(current_chunk)
                        }
                    }
                    chunks.append(chunk_data)
                    chunk_index += 1
                    
                    # Start new chunk with overlap
                    current_chunk = paragraph
                else:
                    # Paragraph too long, split by sentences
                    sentences = paragraph.split('. ')
                    for sentence in sentences:
                        if len(current_chunk) + len(sentence) > self.chunk_size:
                            if current_chunk:
                                chunk_data = {
                                    "chunk_index": chunk_index,
                                    "content": current_chunk.strip(),
                                    "chunk_content": current_chunk.strip(),
                                    "metadata": {
                                        "position": chunk_index,
                                        "word_count": len(current_chunk.split()),
                                        "char_count": len(current_chunk),
                                        "section_path": self._get_section_path(chunk_index, structure),
                                        "token_count": len(current_chunk.split())
                                    },
                                    "contexts": {
                                        "local": self._extract_local_context_enhanced(current_chunk, paragraphs, para_idx),
                                        "document": metadata,
                                        "global": {"document_title": metadata.get("title", "")},
                                        "semantic": await self._extract_semantic_context(current_chunk)
                                    }
                                }
                                chunks.append(chunk_data)
                                chunk_index += 1
                            current_chunk = sentence
                        else:
                            current_chunk += sentence + ". "
            else:
                current_chunk += "\n\n" + paragraph
        
        # Add final chunk
        if current_chunk.strip():
            chunk_data = {
                "chunk_index": chunk_index,
                "content": current_chunk.strip(),
                "chunk_content": current_chunk.strip(),
                "metadata": {
                    "position": chunk_index,
                    "word_count": len(current_chunk.split()),
                    "char_count": len(current_chunk),
                    "section_path": self._get_section_path(chunk_index, structure),
                    "token_count": len(current_chunk.split())
                },
                "contexts": {
                    "local": self._extract_local_context_enhanced(current_chunk, paragraphs, len(paragraphs)-1),
                    "document": metadata,
                    "global": {"document_title": metadata.get("title", "")},
                    "semantic": await self._extract_semantic_context(current_chunk)
                }
            }
            chunks.append(chunk_data)
        
        return chunks
    
    def _get_section_path(self, chunk_index: int, structure: Dict[str, Any]) -> List[str]:
        """Get section path for chunk"""
        # Simplified implementation - would be more sophisticated in practice
        sections = structure.get("sections", [])
        if sections and chunk_index < len(sections):
            return [sections[chunk_index].get("title", f"Section {chunk_index}")]
        return [f"Section {chunk_index}"]
    
    def _extract_local_context_enhanced(self, chunk: str, all_paragraphs: List[str], current_idx: int) -> str:
        """Extract enhanced local context"""
        context_parts = []
        
        # Previous context
        if current_idx > 0:
            context_parts.append(f"Previous: {all_paragraphs[current_idx-1][-100:]}")
        
        # Next context  
        if current_idx < len(all_paragraphs) - 1:
            context_parts.append(f"Next: {all_paragraphs[current_idx+1][:100]}")
        
        return " | ".join(context_parts)
    
    async def _extract_semantic_context(self, chunk: str) -> str:
        """Extract semantic context using AI"""
        try:
            # Simple keyword extraction for now
            words = chunk.split()
            # Get longest words as potential concepts
            concepts = sorted(set(w for w in words if len(w) > 6), key=len, reverse=True)[:5]
            return ", ".join(concepts)
        except Exception as e:
            logger.warning(f"Semantic context extraction failed: {e}")
            return ""