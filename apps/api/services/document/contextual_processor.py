# apps/api/services/document/contextual_processor.py
"""
Enhanced Document Processor with 3-level Contextual Embeddings
Implements intelligent chunking and multi-level context extraction
"""

import asyncio
import hashlib
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import numpy as np
from datetime import datetime

# Document processing
import PyPDF2
import docx
from pdfplumber import PDF
import tiktoken
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

# ML/NLP
from sentence_transformers import SentenceTransformer
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Your imports
from apps.api.core.config import settings
from apps.api.core.database import get_db
from apps.api.models.document import Document, DocumentChunk, ContextualEmbedding

logger = logging.getLogger(__name__)

# Download required NLTK data
def download_nltk_data():
    """Download NLTK data with SSL workaround"""
    import ssl
    import os

    # Create unverified SSL context for NLTK downloads (macOS workaround)
    try:
        _create_unverified_https_context = ssl._create_unverified_context
    except AttributeError:
        # Legacy Python that doesn't verify HTTPS certificates by default
        pass
    else:
        # Handle target environment that doesn't support HTTPS verification
        ssl._create_default_https_context = _create_unverified_https_context

    # Ensure NLTK data directory exists
    nltk_data_dir = os.path.expanduser('~/nltk_data')
    os.makedirs(nltk_data_dir, exist_ok=True)

    # Try to download required packages
    packages = ['punkt', 'punkt_tab', 'stopwords']
    for package in packages:
        try:
            nltk.download(package, quiet=True, raise_on_error=True)
            logger.info(f"Successfully downloaded NLTK package: {package}")
        except Exception as e:
            logger.warning(f"Failed to download NLTK package '{package}': {e}")
            # Check if package already exists locally
            try:
                if package == 'punkt':
                    nltk.data.find('tokenizers/punkt')
                    logger.info(f"NLTK package '{package}' already available locally")
                elif package == 'punkt_tab':
                    nltk.data.find('tokenizers/punkt_tab')
                    logger.info(f"NLTK package '{package}' already available locally")
                elif package == 'stopwords':
                    nltk.data.find('corpora/stopwords')
                    logger.info(f"NLTK package '{package}' already available locally")
            except LookupError:
                logger.warning(f"NLTK package '{package}' not available locally or via download")

# Initialize NLTK data
try:
    download_nltk_data()
except Exception as e:
    logger.warning(f"NLTK initialization failed: {e}. Text processing features may be limited.")
    pass

# Safe NLTK wrappers that fallback when data is unavailable
def safe_sent_tokenize(text: str) -> List[str]:
    """Tokenize text into sentences with fallback"""
    try:
        return sent_tokenize(text)
    except LookupError:
        # Fallback to simple sentence splitting
        logger.debug("NLTK punkt not available, using simple sentence splitting")
        import re
        # Simple sentence splitter based on punctuation
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]

def safe_word_tokenize(text: str) -> List[str]:
    """Tokenize text into words with fallback"""
    try:
        return word_tokenize(text)
    except LookupError:
        # Fallback to simple word splitting
        logger.debug("NLTK punkt not available, using simple word splitting")
        import re
        words = re.findall(r'\b\w+\b', text.lower())
        return words

def safe_get_stopwords() -> set:
    """Get stopwords with fallback"""
    try:
        return set(stopwords.words('english'))
    except LookupError:
        # Fallback to basic stopwords list
        logger.debug("NLTK stopwords not available, using basic stopword list")
        return {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'be',
            'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will',
            'would', 'should', 'could', 'may', 'might', 'must', 'can', 'this',
            'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they'
        }

@dataclass
class ChunkMetadata:
    """Metadata for each chunk"""
    chunk_index: int
    start_char: int
    end_char: int
    start_page: Optional[int] = None
    end_page: Optional[int] = None
    section_path: List[str] = field(default_factory=list)
    heading_level: Optional[int] = None
    is_table: bool = False
    is_list: bool = False
    semantic_type: str = "text"  # text, code, table, list, heading
    entities: List[Dict[str, str]] = field(default_factory=list)
    keywords: List[str] = field(default_factory=list)
    
@dataclass
class ProcessedChunk:
    """Processed chunk with content and metadata"""
    content: str
    metadata: ChunkMetadata
    local_context: Optional[str] = None
    document_context: Optional[str] = None
    global_context: Optional[str] = None
    embedding: Optional[List[float]] = None
    
@dataclass
class DocumentStructure:
    """Document structure information"""
    title: str
    sections: List[Dict[str, Any]]
    hierarchy: Dict[str, List[str]]
    total_pages: int
    total_chunks: int
    document_type: str
    summary: Optional[str] = None
    key_topics: List[str] = field(default_factory=list)

class ModernDocumentProcessor:
    """Enhanced document processor with intelligent chunking and contextual understanding"""
    
    def __init__(self):
        # Initialize NLP models
        self.nlp = spacy.load("en_core_web_sm")
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Chunking parameters
        self.target_chunk_size = 512  # tokens
        self.max_chunk_size = 1024
        self.min_chunk_size = 100
        self.chunk_overlap = 50  # tokens
        
        # Context window sizes
        self.local_context_window = 2  # chunks before/after
        self.document_context_samples = 5  # representative chunks
        self.global_context_docs = 3  # related documents
        
    async def process_document(
        self,
        file_path: str,
        content_type: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Process document with intelligent chunking and structure extraction"""
        try:
            logger.info(f"Processing document: {file_path}")
            
            # Extract raw content and structure
            raw_content, structure_info = await self._extract_content_and_structure(
                file_path, content_type
            )
            
            # Perform intelligent chunking
            chunks = await self._intelligent_chunking(raw_content, structure_info)
            
            # Extract document-level context
            document_structure = await self._analyze_document_structure(
                chunks, structure_info, metadata
            )
            
            # Generate summary
            document_structure.summary = await self._generate_document_summary(chunks)
            
            # Extract key topics
            document_structure.key_topics = await self._extract_key_topics(chunks)
            
            # Calculate file hash
            file_hash = await self._calculate_file_hash(file_path)
            
            return {
                "chunks": chunks,
                "structure": document_structure,
                "metadata": {
                    "file_path": file_path,
                    "content_type": content_type,
                    "file_hash": file_hash,
                    "processed_at": datetime.utcnow().isoformat(),
                    "processor_version": "2.0",
                    **metadata
                } if metadata else {
                    "file_path": file_path,
                    "content_type": content_type,
                    "file_hash": file_hash,
                    "processed_at": datetime.utcnow().isoformat(),
                    "processor_version": "2.0"
                }
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            raise
    
    async def _extract_content_and_structure(
        self,
        file_path: str,
        content_type: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract content and structure based on file type"""
        
        if content_type == "application/pdf":
            return await self._extract_pdf_content(file_path)
        elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            return await self._extract_docx_content(file_path)
        elif content_type.startswith("text/"):
            return await self._extract_text_content(file_path)
        else:
            raise ValueError(f"Unsupported content type: {content_type}")
    
    async def _extract_pdf_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content and structure from PDF"""
        content_parts = []
        structure_info = {
            "pages": [],
            "sections": [],
            "tables": [],
            "images": []
        }

        with open(file_path, 'rb') as f:
            with PDF(f) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text() or ""
                    content_parts.append(text)

                    # Extract structure
                    structure_info["pages"].append({
                        "page_num": page_num + 1,
                        "text": text,
                        "tables": len(page.extract_tables()),
                        "bbox": page.bbox
                    })

                    # Extract tables
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        structure_info["tables"].append({
                            "page": page_num + 1,
                            "index": table_idx,
                            "data": table
                        })

        # Detect sections based on formatting
        full_content = "\n\n".join(content_parts)
        sections = self._detect_sections(full_content)
        structure_info["sections"] = sections

        return full_content, structure_info
    
    async def _extract_docx_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content and structure from DOCX"""
        doc = docx.Document(file_path)
        content_parts = []
        structure_info = {
            "paragraphs": [],
            "sections": [],
            "tables": [],
            "styles": {}
        }
        
        current_section = []
        current_heading = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            style_name = para.style.name
            
            # Detect headings
            if 'Heading' in style_name:
                if current_section:
                    structure_info["sections"].append({
                        "heading": current_heading,
                        "content": "\n".join(current_section),
                        "level": int(style_name[-1]) if style_name[-1].isdigit() else 1
                    })
                current_heading = text
                current_section = []
            else:
                current_section.append(text)
            
            content_parts.append(text)
            structure_info["paragraphs"].append({
                "text": text,
                "style": style_name
            })
        
        # Add last section
        if current_section:
            structure_info["sections"].append({
                "heading": current_heading,
                "content": "\n".join(current_section),
                "level": 1
            })
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            structure_info["tables"].append({
                "index": table_idx,
                "data": table_data
            })
        
        return "\n\n".join(content_parts), structure_info
    
    async def _extract_text_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from plain text file with encoding detection"""
        import chardet

        # Read file in binary mode to detect encoding
        with open(file_path, 'rb') as f:
            raw_data = f.read()

        # Detect encoding
        detected = chardet.detect(raw_data)
        encoding = detected.get('encoding', 'utf-8')
        confidence = detected.get('confidence', 0)

        logger.info(f"Detected encoding: {encoding} (confidence: {confidence:.2f})")

        # Try detected encoding first, fall back to utf-8, then latin-1
        for enc in [encoding, 'utf-8', 'latin-1', 'cp1252']:
            try:
                content = raw_data.decode(enc)
                logger.info(f"Successfully decoded with {enc}")
                break
            except (UnicodeDecodeError, AttributeError, LookupError):
                continue
        else:
            # Last resort: decode with errors='ignore'
            content = raw_data.decode('utf-8', errors='ignore')
            logger.warning(f"Used UTF-8 with errors='ignore' for {file_path}")

        structure_info = {
            "sections": self._detect_sections(content),
            "paragraphs": content.split('\n\n'),
            "detected_encoding": encoding,
            "encoding_confidence": confidence
        }

        return content, structure_info
    
    def _detect_sections(self, text: str) -> List[Dict[str, Any]]:
        """Detect sections in text based on patterns"""
        sections = []
        lines = text.split('\n')
        
        current_section = {
            "heading": "Introduction",
            "content": [],
            "level": 1,
            "start_line": 0
        }
        
        # Common heading patterns
        heading_patterns = [
            (r'^#{1,6}\s+(.+)$', 'markdown'),  # Markdown headings
            (r'^([A-Z][A-Z\s]+)$', 'uppercase'),  # UPPERCASE HEADINGS
            (r'^(\d+\.?\s+.+)$', 'numbered'),  # 1. Heading
            (r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*):?\s*$', 'title_case'),  # Title Case
        ]
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue
                
            # Check if line is a heading
            is_heading = False
            for pattern, style in heading_patterns:
                import re
                match = re.match(pattern, stripped)
                if match:
                    # Save current section
                    if current_section["content"]:
                        sections.append(current_section)
                    
                    # Start new section
                    level = stripped.count('#') if style == 'markdown' else 1
                    heading_text = match.group(1).strip('#').strip()
                    
                    current_section = {
                        "heading": heading_text,
                        "content": [],
                        "level": level,
                        "start_line": i
                    }
                    is_heading = True
                    break
            
            if not is_heading:
                current_section["content"].append(line)
        
        # Add last section
        if current_section["content"]:
            sections.append(current_section)
        
        # Clean up sections
        for section in sections:
            section["content"] = "\n".join(section["content"]).strip()
        
        return sections
    
    async def _intelligent_chunking(
        self,
        content: str,
        structure_info: Dict[str, Any]
    ) -> List[ProcessedChunk]:
        """Perform intelligent chunking that preserves semantic boundaries"""
        chunks = []
        
        # Get sections
        sections = structure_info.get("sections", [])
        if not sections:
            # Fallback to paragraph-based chunking
            sections = [{"heading": "Document", "content": content, "level": 1}]
        
        chunk_index = 0
        
        for section in sections:
            section_content = section["content"]
            section_path = [section["heading"]]
            
            # Skip empty sections
            if not section_content.strip():
                continue
            
            # Chunk the section
            section_chunks = await self._chunk_section(
                section_content,
                section_path,
                section.get("level", 1),
                chunk_index
            )
            
            chunks.extend(section_chunks)
            chunk_index += len(section_chunks)
        
        return chunks
    
    async def _chunk_section(
        self,
        content: str,
        section_path: List[str],
        heading_level: int,
        start_index: int
    ) -> List[ProcessedChunk]:
        """Chunk a section intelligently"""
        chunks = []

        # Split into sentences using safe wrapper
        sentences = safe_sent_tokenize(content)
        
        current_chunk = []
        current_tokens = 0
        chunk_start_char = 0
        
        for sent_idx, sentence in enumerate(sentences):
            sent_tokens = len(self.tokenizer.encode(sentence))
            
            # Check if adding this sentence would exceed target size
            if current_tokens + sent_tokens > self.target_chunk_size and current_chunk:
                # Create chunk
                chunk_content = " ".join(current_chunk)
                
                # Detect semantic type
                semantic_type = self._detect_semantic_type(chunk_content)
                
                # Extract entities and keywords
                entities = self._extract_entities(chunk_content)
                keywords = self._extract_keywords(chunk_content)
                
                chunk = ProcessedChunk(
                    content=chunk_content,
                    metadata=ChunkMetadata(
                        chunk_index=start_index + len(chunks),
                        start_char=chunk_start_char,
                        end_char=chunk_start_char + len(chunk_content),
                        section_path=section_path,
                        heading_level=heading_level,
                        semantic_type=semantic_type,
                        entities=entities,
                        keywords=keywords
                    )
                )
                
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_sentences = self._get_overlap_sentences(current_chunk, self.chunk_overlap)
                current_chunk = overlap_sentences + [sentence]
                current_tokens = sum(len(self.tokenizer.encode(s)) for s in current_chunk)
                chunk_start_char += len(chunk_content) - len(" ".join(overlap_sentences))
            else:
                current_chunk.append(sentence)
                current_tokens += sent_tokens
        
        # Add final chunk
        if current_chunk:
            chunk_content = " ".join(current_chunk)
            
            chunk = ProcessedChunk(
                content=chunk_content,
                metadata=ChunkMetadata(
                    chunk_index=start_index + len(chunks),
                    start_char=chunk_start_char,
                    end_char=chunk_start_char + len(chunk_content),
                    section_path=section_path,
                    heading_level=heading_level,
                    semantic_type=self._detect_semantic_type(chunk_content),
                    entities=self._extract_entities(chunk_content),
                    keywords=self._extract_keywords(chunk_content)
                )
            )
            
            chunks.append(chunk)
        
        return chunks
    
    def _get_overlap_sentences(self, sentences: List[str], target_tokens: int) -> List[str]:
        """Get sentences for overlap"""
        overlap_sentences = []
        tokens = 0
        
        for sent in reversed(sentences):
            sent_tokens = len(self.tokenizer.encode(sent))
            if tokens + sent_tokens <= target_tokens:
                overlap_sentences.insert(0, sent)
                tokens += sent_tokens
            else:
                break
        
        return overlap_sentences
    
    def _detect_semantic_type(self, text: str) -> str:
        """Detect the semantic type of text"""
        # Simple heuristics - can be enhanced with ML
        if text.count('|') > 5 or text.count('\t') > 5:
            return "table"
        elif any(text.startswith(marker) for marker in ['•', '-', '*', '1.', '2.']):
            return "list"
        elif text.count('```') >= 2 or text.count('    ') > len(text.split('\n')) / 2:
            return "code"
        elif len(text.split()) < 10 and text.isupper():
            return "heading"
        else:
            return "text"
    
    def _extract_entities(self, text: str) -> List[Dict[str, str]]:
        """Extract named entities from text"""
        doc = self.nlp(text)
        entities = []
        
        for ent in doc.ents:
            entities.append({
                "text": ent.text,
                "type": ent.label_,
                "start": ent.start_char,
                "end": ent.end_char
            })
        
        return entities
    
    def _extract_keywords(self, text: str, max_keywords: int = 5) -> List[str]:
        """Extract keywords using TF-IDF"""
        try:
            # Tokenize and remove stopwords using safe wrappers
            stop_words = safe_get_stopwords()
            words = safe_word_tokenize(text.lower())
            words = [w for w in words if w.isalnum() and w not in stop_words and len(w) > 2]
            
            if not words:
                return []
            
            # Simple frequency-based extraction
            from collections import Counter
            word_freq = Counter(words)
            
            # Get top keywords
            keywords = [word for word, _ in word_freq.most_common(max_keywords)]
            
            return keywords
        except Exception as e:
            logger.error(f"Keyword extraction failed: {e}")
            return []
    
    async def _analyze_document_structure(
        self,
        chunks: List[ProcessedChunk],
        structure_info: Dict[str, Any],
        metadata: Optional[Dict[str, Any]]
    ) -> DocumentStructure:
        """Analyze overall document structure"""
        
        # Build hierarchy
        hierarchy = {}
        for chunk in chunks:
            path = chunk.metadata.section_path
            if path:
                section = path[0]
                if section not in hierarchy:
                    hierarchy[section] = []
                hierarchy[section].append(chunk.metadata.chunk_index)
        
        # Get document type
        doc_type = "unknown"
        if metadata:
            doc_type = metadata.get("document_type", "unknown")
        elif "contract" in str(chunks).lower():
            doc_type = "contract"
        elif "report" in str(chunks).lower():
            doc_type = "report"
        
        return DocumentStructure(
            title=metadata.get("title", "Untitled") if metadata else "Untitled",
            sections=structure_info.get("sections", []),
            hierarchy=hierarchy,
            total_pages=len(structure_info.get("pages", [])),
            total_chunks=len(chunks),
            document_type=doc_type
        )
    
    async def _generate_document_summary(self, chunks: List[ProcessedChunk]) -> str:
        """Generate document summary from chunks"""
        # Take first few and last few chunks
        sample_chunks = chunks[:3] + chunks[-2:] if len(chunks) > 5 else chunks
        
        # Combine content
        combined_text = " ".join([chunk.content for chunk in sample_chunks])

        # Simple extractive summary (can be replaced with abstractive)
        sentences = safe_sent_tokenize(combined_text)
        if len(sentences) > 5:
            # Use first 2 and last sentence
            summary = " ".join(sentences[:2] + [sentences[-1]])
        else:
            summary = combined_text
        
        # Limit length
        if len(summary.split()) > 100:
            words = summary.split()[:100]
            summary = " ".join(words) + "..."
        
        return summary
    
    async def _extract_key_topics(self, chunks: List[ProcessedChunk]) -> List[str]:
        """Extract key topics from document"""
        # Collect all keywords
        all_keywords = []
        for chunk in chunks:
            all_keywords.extend(chunk.metadata.keywords)
        
        # Count frequency
        from collections import Counter
        keyword_freq = Counter(all_keywords)
        
        # Get top topics
        topics = [topic for topic, _ in keyword_freq.most_common(10)]
        
        return topics
    
    async def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        sha256_hash = hashlib.sha256()
        
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        
        return sha256_hash.hexdigest()

class ContextualEmbeddingGenerator:
    """Generate contextual embeddings with 3-level context"""
    
    def __init__(self):
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.context_weights = {
            "local": 0.5,
            "document": 0.3,
            "global": 0.2
        }
    
    async def generate_embeddings(
        self,
        chunks: List[ProcessedChunk],
        document_metadata: Dict[str, Any],
        use_voyage: bool = False
    ) -> List[Dict[str, Any]]:
        """Generate contextual embeddings for chunks"""
        embeddings = []
        
        for i, chunk in enumerate(chunks):
            # Get local context
            local_context = await self._get_local_context(chunks, i)
            
            # Get document context
            document_context = await self._get_document_context(
                chunks,
                document_metadata
            )
            
            # Get global context (would query existing documents)
            global_context = await self._get_global_context(
                chunk,
                document_metadata
            )
            
            # Generate contextual embedding
            embedding = await self._generate_contextual_embedding(
                chunk.content,
                local_context,
                document_context,
                global_context,
                use_voyage
            )
            
            embeddings.append({
                "chunk_id": f"{document_metadata.get('document_id', 'unknown')}_chunk_{i}",
                "content": chunk.content,
                "embedding": embedding,
                "metadata": {
                    **chunk.metadata.__dict__,
                    "local_context": local_context,
                    "document_context": document_context,
                    "global_context": global_context
                }
            })
        
        return embeddings
    
    async def _get_local_context(
        self,
        chunks: List[ProcessedChunk],
        chunk_index: int,
        window_size: int = 2
    ) -> str:
        """Get local context from surrounding chunks"""
        context_parts = []
        
        # Previous chunks
        for i in range(max(0, chunk_index - window_size), chunk_index):
            context_parts.append(f"[Previous {chunk_index - i}]: {chunks[i].content[:100]}...")
        
        # Current chunk section info
        current_chunk = chunks[chunk_index]
        if current_chunk.metadata.section_path:
            context_parts.append(f"[Section]: {' > '.join(current_chunk.metadata.section_path)}")
        
        # Next chunks
        for i in range(chunk_index + 1, min(len(chunks), chunk_index + window_size + 1)):
            context_parts.append(f"[Next {i - chunk_index}]: {chunks[i].content[:100]}...")
        
        return " ".join(context_parts)
    
    async def _get_document_context(
        self,
        chunks: List[ProcessedChunk],
        metadata: Dict[str, Any]
    ) -> str:
        """Get document-level context"""
        context_parts = []
        
        # Document metadata
        if metadata.get("title"):
            context_parts.append(f"[Document]: {metadata['title']}")
        
        if metadata.get("document_type"):
            context_parts.append(f"[Type]: {metadata['document_type']}")
        
        # Section overview
        sections = set()
        for chunk in chunks:
            if chunk.metadata.section_path:
                sections.add(chunk.metadata.section_path[0])
        
        if sections:
            context_parts.append(f"[Sections]: {', '.join(list(sections)[:5])}")
        
        # Key topics
        all_keywords = []
        for chunk in chunks:
            all_keywords.extend(chunk.metadata.keywords)
        
        from collections import Counter
        top_keywords = [kw for kw, _ in Counter(all_keywords).most_common(5)]
        if top_keywords:
            context_parts.append(f"[Topics]: {', '.join(top_keywords)}")
        
        return " ".join(context_parts)
    
    async def _get_global_context(
        self,
        chunk: ProcessedChunk,
        metadata: Dict[str, Any]
    ) -> str:
        """Get global context from related documents"""
        # In production, this would query the vector store for related documents
        # For now, return placeholder
        context_parts = []
        
        # Document collection info
        if metadata.get("collection"):
            context_parts.append(f"[Collection]: {metadata['collection']}")
        
        # Related topics (would come from knowledge graph)
        if chunk.metadata.keywords:
            context_parts.append(f"[Related Topics]: {', '.join(chunk.metadata.keywords[:3])}")
        
        return " ".join(context_parts) if context_parts else "[Global Context]: No related documents found"
    
    async def _generate_contextual_embedding(
        self,
        content: str,
        local_context: str,
        document_context: str,
        global_context: str,
        use_voyage: bool = False
    ) -> List[float]:
        """Generate embedding with weighted context"""
        
        # Combine content with contexts
        combined_text = f"{content} {local_context} {document_context} {global_context}"
        
        if use_voyage and hasattr(settings, 'voyage_api_key'):
            # Use Voyage AI for embeddings
            try:
                import voyageai
                voyage_client = voyageai.Client(api_key=settings.voyage_api_key)
                result = voyage_client.embed(
                    [combined_text],
                    model="voyage-2"
                )
                embedding = result.embeddings[0]
            except Exception as e:
                logger.error(f"Voyage AI embedding failed: {e}")
                # Fallback to local model
                embedding = self.embedding_model.encode(combined_text).tolist()
        else:
            # Use local sentence transformer
            embedding = self.embedding_model.encode(combined_text).tolist()
        
        return embedding
    
    async def generate_contextual_embedding(
        self,
        text: str,
        context_level: str = "local",
        conversation_history: Optional[List[Dict]] = None,
        document_context: Optional[List[str]] = None
    ) -> List[float]:
        """Generate a single contextual embedding for queries"""
        
        context_parts = [text]
        
        # Add conversation context
        if conversation_history:
            recent_messages = conversation_history[-3:]  # Last 3 messages
            for msg in recent_messages:
                context_parts.append(f"[{msg['role']}]: {msg['content'][:100]}")
        
        # Add document context
        if document_context:
            context_parts.append(f"[Documents]: {', '.join(document_context[:3])}")
        
        # Add context level indicator
        context_parts.append(f"[Level]: {context_level}")
        
        combined_text = " ".join(context_parts)
        
        # Generate embedding
        embedding = self.embedding_model.encode(combined_text).tolist()
        
        return embedding